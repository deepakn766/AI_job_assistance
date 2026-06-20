"""
agents/resume_optimizer_agent.py
---------------------------------
AGENT 4: ATS Resume Optimization Agent

FIXED VERSION — changes from original:
  1. Upgraded model llama3-8b-8192 -> llama-3.3-70b-versatile. The 8B
     model is weak at following the "never fabricate, rewrite everything
     faithfully" instruction set and tends to drop sections under
     pressure to stay concise — switching to the same 70B model used
     for parsing makes output quality consistent and far more reliable
     for a full-resume rewrite.
  2. max_tokens raised 1800 -> 3000, and response_format=json_object
     added, for the same truncation/parsing reasons as the resume agent.
  3. experience_str / projects_str / education_str no longer silently
     slice to short previews — full lists are passed in (still capped
     generously) so nothing is dropped before it even reaches the model.
  4. Added an explicit guard: if profile has NO experience/projects at
     all, the prompt tells Groq this clearly instead of just leaving it
     blank, AND the function prints a warning so you know to look
     upstream (in the Resume Agent / parsing step) rather than assuming
     this agent ate your data.
  5. Raw response is logged on JSON failure, same as the resume agent.
"""

import os
import json
import re
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

GENERATED_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "generated")
os.makedirs(GENERATED_DIR, exist_ok=True)


def get_groq_client() -> Groq:
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY not set in .env")
    return Groq(api_key=api_key)


# ─── Core Optimization Function ───────────────────────────────────────────────

def optimize_resume(profile: dict, job: dict) -> dict:
    """
    Main entry point for the Resume Optimizer Agent.

    Flow:
    1. Build a targeted prompt with profile + JD
    2. Ask Groq to optimize each resume section
    3. Return structured optimized resume dict
    4. Also save as DOCX file

    Args:
        profile: Candidate profile dict from Resume Agent
        job:     Selected job dict (title, company, description)

    Returns:
        optimized dict with sections + DOCX file path
    """
    client = get_groq_client()

    skills_list = profile.get("skills", [])
    experience_list = profile.get("experience", [])
    projects_list = profile.get("projects", [])
    education_list = profile.get("education", [])

    # FIX: surface this immediately instead of silently generating a
    # thin resume. If you see this warning, the problem is upstream in
    # resume_agent.py / analyze_resume(), not in this file.
    if not experience_list and not projects_list:
        print("[Optimizer Agent] WARNING: profile has NO experience and NO projects. "
              "The optimized resume will only contain summary/skills/education. "
              "Check that analyze_resume() actually extracted these sections "
              "from the original resume text.")

    skills_str = ", ".join(skills_list[:30])
    experience_str = "\n".join([f"- {exp}" for exp in experience_list[:10]])  # FIX: 5 -> 10
    projects_str = "\n".join([f"- {proj}" for proj in projects_list[:8]])      # FIX: 4 -> 8
    education_str = "\n".join([f"- {edu}" for edu in education_list[:5]])      # FIX: 3 -> 5

    prompt = f"""
You are an expert ATS resume writer and career coach.

Your task: Optimize this candidate's ENTIRE resume for the specific job below.
You must produce a COMPLETE resume — summary, skills, experience, projects,
and education — not just a subset. Do not drop any section that has content
in the candidate profile below.

STRICT RULES (NEVER BREAK THESE):
1. Do NOT invent, fabricate, or add fake experience, skills, or projects.
2. Only IMPROVE wording, professionalism, and keyword alignment.
3. Use strong action verbs (Developed, Built, Optimized, Implemented, etc.)
4. Mirror keywords from the job description naturally.
5. Keep bullet points concise: max 2 lines each.
6. Rewrite EVERY experience entry and EVERY project listed below — do not
   skip any of them, and do not collapse multiple entries into one.
7. If a section below is genuinely empty, leave the corresponding output
   array empty — do not invent content to fill it.

CANDIDATE PROFILE:
Name: {profile.get('name', '')}
Current Skills: {skills_str}

Experience ({len(experience_list)} entries — rewrite ALL of them):
{experience_str if experience_str else "(none provided)"}

Projects ({len(projects_list)} entries — rewrite ALL of them):
{projects_str if projects_str else "(none provided)"}

Education:
{education_str if education_str else "(none provided)"}

TARGET JOB:
Title: {job.get('title', '')}
Company: {job.get('company', '')}
Job Description:
{job.get('description', '')[:1500]}

OUTPUT: Return ONLY a valid JSON object with these keys:
{{
  "summary": "3-4 sentence ATS-optimized professional summary tailored to this role",
  "skills_section": ["skill1", "skill2", ...],
  "experience_bullets": ["Improved bullet 1", "Improved bullet 2", ...],
  "projects_bullets": ["Project bullet 1", "Project bullet 2", ...],
  "keywords_added": ["keyword1", "keyword2", ...]
}}

The experience_bullets and projects_bullets arrays must collectively cover
EVERY entry from the candidate profile above — do not output fewer bullets
than there are entries unless an entry genuinely had nothing to rewrite.
No markdown, no extra text. Valid JSON only.
"""

    optimized = None
    try:
        response = client.chat.completions.create(
            # FIX: upgraded from llama3-8b-8192. The 8B model was prone to
            # truncating/dropping sections on long, multi-section rewrites.
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=3000,  # FIX: was 1800 — too low for a full resume rewrite
            response_format={"type": "json_object"},  # FIX: forces valid JSON
        )
        text = response.choices[0].message.content.strip()
        text = re.sub(r"```json|```", "", text).strip()
        optimized = json.loads(text)

    except json.JSONDecodeError as e:
        print(f"[Optimizer Agent] JSON parsing failed: {e}")
        print(f"[Optimizer Agent] Raw response that failed to parse:\n{text[:2000]}")
        optimized = None

    except Exception as e:
        print(f"[Optimizer Agent] Groq API call failed: {e}")
        optimized = None

    if optimized is None:
        print("[Optimizer Agent] Using fallback (lightly formatted original data, no AI rewrite).")
        optimized = {
            "summary": profile.get("summary", "Experienced professional seeking new opportunities."),
            "skills_section": skills_list[:20],
            "experience_bullets": [f"{exp}" for exp in experience_list],
            "projects_bullets": [f"{proj}" for proj in projects_list],
            "keywords_added": []
        }

    # Add metadata
    optimized["candidate_name"] = profile.get("name", "Candidate")
    optimized["candidate_email"] = profile.get("email", "")
    optimized["target_job"] = f"{job.get('title', '')} at {job.get('company', '')}"
    optimized["education"] = education_list

    # FIX: final safety net — if the AI rewrite somehow returned fewer
    # experience/project bullets than the candidate actually has, fall
    # back to the original (un-optimized but COMPLETE) entries for the
    # missing ones rather than silently truncating the resume.
    if experience_list and len(optimized.get("experience_bullets", [])) < len(experience_list):
        print(f"[Optimizer Agent] WARNING: model returned "
              f"{len(optimized.get('experience_bullets', []))} experience bullets but profile has "
              f"{len(experience_list)} entries. Falling back to original experience text to avoid data loss.")
        optimized["experience_bullets"] = [f"{exp}" for exp in experience_list]

    if projects_list and len(optimized.get("projects_bullets", [])) < len(projects_list):
        print(f"[Optimizer Agent] WARNING: model returned "
              f"{len(optimized.get('projects_bullets', []))} project bullets but profile has "
              f"{len(projects_list)} entries. Falling back to original project text to avoid data loss.")
        optimized["projects_bullets"] = [f"{proj}" for proj in projects_list]

    # Save DOCX
    docx_path = save_as_docx(optimized, profile, job)
    optimized["docx_path"] = docx_path

    return optimized


# ─── DOCX Export ──────────────────────────────────────────────────────────────

def save_as_docx(optimized: dict, profile: dict, job: dict) -> str:
    """
    Create a clean, ATS-friendly DOCX resume from the optimized data.

    ATS-friendly design principles:
    - No tables, no text boxes (ATS can't read them)
    - Standard fonts (Calibri)
    - Clear section headers
    - Simple formatting
    """
    doc = Document()

    # Set document margins
    from docx.shared import Inches
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    def add_heading(text, level=1):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 11)
        if level == 2:
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)  # Professional blue
        para.space_after = Pt(4)
        return para

    def add_body(text, bold=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.bold = bold
        para.space_after = Pt(2)
        return para

    def add_bullet(text):
        # FIX: use Word's actual bullet list style instead of relying on
        # the LLM to prepend a "•" character (which it sometimes drops).
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(text)
        run.font.size = Pt(10)
        para.space_after = Pt(2)
        return para

    # --- Header: Name + Contact ---
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(optimized.get("candidate_name", "Candidate"))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email = optimized.get("candidate_email", "")
    phone = profile.get("phone", "")
    contact_text = " | ".join(filter(None, [email, phone]))
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(10)

    doc.add_paragraph()  # Spacer

    # --- Professional Summary ---
    add_heading("PROFESSIONAL SUMMARY", level=2)
    add_body(optimized.get("summary", ""))
    doc.add_paragraph()

    # --- Skills ---
    add_heading("TECHNICAL SKILLS", level=2)
    skills = optimized.get("skills_section", [])
    if skills:
        add_body("  •  ".join(skills[:25]))  # FIX: 20 -> 25
    doc.add_paragraph()

    # --- Experience ---
    if optimized.get("experience_bullets"):
        add_heading("EXPERIENCE", level=2)
        for bullet_text in optimized["experience_bullets"]:
            add_bullet(bullet_text.lstrip("•- ").strip())
        doc.add_paragraph()

    # --- Projects ---
    if optimized.get("projects_bullets"):
        add_heading("PROJECTS", level=2)
        for bullet_text in optimized["projects_bullets"]:
            add_bullet(bullet_text.lstrip("•- ").strip())
        doc.add_paragraph()

    # --- Education ---
    add_heading("EDUCATION", level=2)
    for edu in optimized.get("education", []):
        add_body(edu)

    # --- ATS Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Optimized for: {optimized.get('target_job', '')} | Generated by Job Assistant"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save file
    safe_name = re.sub(r"[^\w]", "_", optimized.get("candidate_name", "resume"))
    file_path = os.path.join(GENERATED_DIR, f"optimized_resume_{safe_name}.docx")
    doc.save(file_path)
    print(f"[Optimizer Agent] DOCX saved: {file_path}")
    return file_path